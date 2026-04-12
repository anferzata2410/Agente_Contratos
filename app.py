"""
Aplicacion web - Formulario Check_List para Agente de Contratos
Genera el documento "Contrato Mutuo Solo Interes" como .docx descargable.
"""

import re
import json
import shutil
from pathlib import Path
from datetime import datetime, date

from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from num2words import num2words

CHECKLIST_PATH = Path(__file__).parent / "Documentacion" / "Check_List.docx"

app = Flask(__name__)

BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"
DATA_DIR.mkdir(exist_ok=True)
TEMPLATE_PATH = BASE_DIR / "Documentacion" / "Contrato mutuo solo interés.docx"
OUTPUT_DIR = BASE_DIR / "Contratos_Generados"
OUTPUT_DIR.mkdir(exist_ok=True)


# ═══════════════════════════════════════════════
# UTILIDADES PARA GENERAR TEXTO LEGAL
# ═══════════════════════════════════════════════

def numero_a_texto(n: int) -> str:
    """Convierte un numero a texto en espanol. Ej: 90000000 -> 'noventa millones'"""
    return num2words(n, lang="es")


def monto_a_texto_legal(n: int) -> str:
    """
    Genera formato legal completo.
    Ej: 90000000 -> 'NOVENTA MILLONES DE PESOS MONEDA CORRIENTE (COP$90.000.000)'
    """
    texto = numero_a_texto(n).upper()
    formato = f"{n:,.0f}".replace(",", ".")
    return f"{texto} DE PESOS MONEDA CORRIENTE (COP${formato})"


def monto_a_texto_legal_min(n: int) -> str:
    """
    Genera formato legal en minusculas.
    Ej: 1620000 -> 'Un millon seiscientos veinte mil pesos moneda corriente (COP$1.620.000)'
    """
    texto = numero_a_texto(n).capitalize()
    formato = f"{n:,.0f}".replace(",", ".")
    return f"{texto} pesos moneda corriente (COP${formato})"


def formato_pesos(n: int) -> str:
    """Formatea un numero como pesos: $90.000.000"""
    return f"${n:,.0f}".replace(",", ".")


def formato_pesos_sin_signo(n: int) -> str:
    """Formatea un numero sin signo: 90.000.000"""
    return f"{n:,.0f}".replace(",", ".")


def limpiar_monto(valor: str) -> int:
    """Convierte '180.000.000' o '180000000' a int."""
    if not valor:
        return 0
    limpio = valor.replace("$", "").replace(".", "").replace(",", "").replace(" ", "").strip()
    try:
        return int(limpio)
    except ValueError:
        return 0


def fecha_a_texto_legal(d: date) -> str:
    """Convierte una fecha a texto legal. Ej: 2026-04-08 -> 'Ocho (8) de abril de dos mil veintiseis (2026)'"""
    meses = ["", "enero", "febrero", "marzo", "abril", "mayo", "junio",
             "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
    dia_texto = numero_a_texto(d.day).capitalize()
    anio_texto = numero_a_texto(d.year)
    return f"{dia_texto} ({d.day}) de {meses[d.month]} de {anio_texto} ({d.year})"


def plazo_a_texto(meses: int) -> str:
    """Ej: 60 -> 'Sesenta (60) meses'"""
    texto = numero_a_texto(meses).capitalize()
    return f"{texto} ({meses}) meses"


# ═══════════════════════════════════════════════
# MANIPULACION DE DOCUMENTOS WORD
# ═══════════════════════════════════════════════

def reemplazar_en_parrafo(parrafo, buscar: str, reemplazar: str) -> bool:
    texto_completo = parrafo.text
    if buscar not in texto_completo:
        return False
    for run in parrafo.runs:
        if buscar in run.text:
            run.text = run.text.replace(buscar, reemplazar)
            return True
    nuevo_texto = texto_completo.replace(buscar, reemplazar)
    if parrafo.runs:
        parrafo.runs[0].text = nuevo_texto
        for run in parrafo.runs[1:]:
            run.text = ""
        return True
    return False


def reemplazar_en_documento(doc: Document, buscar: str, reemplazar: str) -> int:
    count = 0
    for para in doc.paragraphs:
        if reemplazar_en_parrafo(para, buscar, reemplazar):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if reemplazar_en_parrafo(para, buscar, reemplazar):
                        count += 1
    return count


# ═══════════════════════════════════════════════
# ENRIQUECIMIENTO DE DATOS DEL FORMULARIO
# ═══════════════════════════════════════════════

def enriquecer_datos(datos_form: dict) -> dict:
    """
    Toma los datos crudos del formulario web y genera la estructura
    enriquecida que necesita el generador de contratos.
    """
    # Soportar lista de deudores o deudor individual
    deudores_raw = datos_form.get("deudores", [])
    if not deudores_raw:
        d = datos_form.get("deudor", {})
        deudores_raw = [d] if d.get("nombre") else []
    acreedores_raw = datos_form.get("acreedores", [])
    inmueble = datos_form.get("inmueble", {})
    prestamo = datos_form.get("prestamo", {})

    # Monto total del credito = suma de participaciones de TODOS los deudores
    # Fallback al campo "monto" del prestamo si no hay participaciones
    monto_deudores = sum(limpiar_monto(d.get("participacion_monto", "0")) for d in deudores_raw)
    monto_prestamo = limpiar_monto(prestamo.get("monto", "0"))
    monto_total = monto_deudores if monto_deudores > 0 else monto_prestamo

    plazo = int(prestamo.get("plazo_meses", "60") or "60")
    tasa = prestamo.get("tasa_mensual", "1.80%")
    cuota_total = limpiar_monto(prestamo.get("cuota_mensual", "0"))
    comision_total = limpiar_monto(prestamo.get("comision_aluri", "0"))
    num_acreedores = len(acreedores_raw) if acreedores_raw else 1

    hoy = date.today()
    fecha_texto = fecha_a_texto_legal(hoy)

    # Enriquecer acreedores
    acreedores = []
    for acr in acreedores_raw:
        monto_part = limpiar_monto(acr.get("participacion_monto", "0"))
        # Cuota y comision proporcional a la participacion de cada acreedor
        if monto_total > 0 and monto_part > 0:
            proporcion = monto_part / monto_total
            cuota_ind = round(cuota_total * proporcion)
            comision_ind = round(comision_total * proporcion)
        else:
            cuota_ind = int(cuota_total / num_acreedores) if num_acreedores > 0 else cuota_total
            comision_ind = int(comision_total / num_acreedores) if num_acreedores > 0 else comision_total
        monto_inicial_ind = monto_part * 35 // 100
        monto_restante_ind = monto_part - monto_inicial_ind

        pct_raw = acr.get("participacion_porcentaje", "")

        acreedores.append({
            "nombre_completo": acr.get("nombre", "").upper(),
            "nombre_completo_mayuscula": acr.get("nombre", "").upper(),
            "cc": acr.get("cc", ""),
            "cc_expedicion": acr.get("cc_expedicion", ""),
            "direccion": acr.get("direccion", ""),
            "email": acr.get("email", ""),
            "telefono": acr.get("telefono", ""),
            "estado_civil": acr.get("estado_civil", ""),
            "participacion_porcentaje": pct_raw,
            "participacion_monto": monto_part,
            "participacion_texto": monto_a_texto_legal(monto_part),
            "cuenta_bancaria": acr.get("cuenta_bancaria", ""),
            "cuota_mensual_individual": cuota_ind,
            "cuota_mensual_texto": monto_a_texto_legal_min(cuota_ind),
            "comision_aluri_individual": comision_ind,
            "monto_inicial": monto_inicial_ind,
            "monto_restante": monto_restante_ind,
        })

    monto_inicial_total = monto_total * 35 // 100
    monto_restante_total = monto_total - monto_inicial_total

    return {
        "fecha_firma": fecha_texto,
        "deudores": [
            {
                "nombre_completo": d.get("nombre", "").upper(),
                "nombre_completo_mayuscula": d.get("nombre", "").upper(),
                "cc": d.get("cc", ""),
                "cc_expedicion": d.get("cc_expedicion", ""),
                "direccion": d.get("direccion", ""),
                "email": d.get("email", ""),
                "telefono": d.get("telefono", ""),
                "estado_civil": d.get("estado_civil", ""),
                "municipio": d.get("cc_expedicion", ""),
            }
            for d in deudores_raw
            if d.get("nombre", "").strip()
        ],
        "codeudores": [
            {
                "nombre_completo": c.get("nombre", "").upper(),
                "nombre_completo_mayuscula": c.get("nombre", "").upper(),
                "cc": c.get("cc", ""),
                "cc_expedicion": c.get("cc_expedicion", ""),
                "direccion": c.get("direccion", ""),
                "email": c.get("email", ""),
                "telefono": c.get("telefono", ""),
                "estado_civil": c.get("estado_civil", ""),
            }
            for c in datos_form.get("codeudores", [])
            if c.get("nombre", "").strip()
        ],
        "acreedores": acreedores,
        "inmueble": {
            "matricula_inmobiliaria": inmueble.get("matricula_inmobiliaria", ""),
            "oficina_registro": "",
            "cedula_catastral": inmueble.get("cedula_catastral", ""),
            "chip": inmueble.get("chip", ""),
            "direccion_corta": inmueble.get("direccion", ""),
            "descripcion_completa": inmueble.get("descripcion", ""),
            "linderos": inmueble.get("linderos", ""),
        },
        "prestamo": {
            "monto_total": monto_total,
            "monto_total_texto": monto_a_texto_legal(monto_total),
            "monto_inicial_credito": monto_inicial_total,
            "monto_inicial_texto": monto_a_texto_legal(monto_inicial_total),
            "monto_restante": monto_restante_total,
            "monto_restante_texto": monto_a_texto_legal(monto_restante_total),
            "plazo_meses": plazo,
            "plazo_texto": plazo_a_texto(plazo),
            "tasa_mensual": tasa,
            "cuota_mensual_total": cuota_total,
            "cuota_mensual_total_texto": monto_a_texto_legal_min(cuota_total),
            "comision_aluri_total": comision_total,
            "comision_aluri_por_acreedor": int(comision_total / num_acreedores) if num_acreedores else comision_total,
            "servicios_aluri_texto": formato_pesos(int(comision_total / num_acreedores) if num_acreedores else comision_total),
        },
    }


# ═══════════════════════════════════════════════
# GENERAR CONTRATO DESDE FORMULARIO
# ═══════════════════════════════════════════════

def insertar_fila_tabla(tabla, despues_de_fila, col0_texto, col1_texto, negrita_col1=False):
    """
    Inserta una fila nueva en una tabla DESPUES de la fila indicada.
    Copia el formato de la fila de referencia.
    """
    from copy import deepcopy

    fila_ref = tabla.rows[despues_de_fila]
    nueva_tr = deepcopy(fila_ref._tr)

    # Limpiar contenido de la nueva fila
    ns_w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for tc in nueva_tr.findall(f'.//{ns_w}tc'):
        for p in tc.findall(f'{ns_w}p'):
            for r in p.findall(f'{ns_w}r'):
                for t in r.findall(f'{ns_w}t'):
                    t.text = ""

    # Insertar despues de la fila de referencia
    fila_ref._tr.addnext(nueva_tr)

    # Escribir textos en la nueva fila (ahora es despues_de_fila + 1)
    nueva_fila = tabla.rows[despues_de_fila + 1]
    # Col 0
    p0 = nueva_fila.cells[0].paragraphs[0]
    if p0.runs:
        p0.runs[0].text = col0_texto
    else:
        p0.add_run(col0_texto)
    # Col 1
    if len(nueva_fila.cells) > 1:
        p1 = nueva_fila.cells[1].paragraphs[0]
        if p1.runs:
            p1.runs[0].text = col1_texto
            if negrita_col1:
                p1.runs[0].bold = True
        else:
            run = p1.add_run(col1_texto)
            if negrita_col1:
                run.bold = True


def escribir_celda(tabla, fila, col, texto, negrita=False):
    """Escribe texto en una celda, limpiando TODO el contenido previo (runs, hyperlinks, etc.)."""
    celda = tabla.rows[fila].cells[col]
    ns_w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    # Limpiar todos los párrafos: runs, hyperlinks y cualquier otro contenido
    for p in celda.paragraphs:
        for run in p.runs:
            run.text = ""
        # Eliminar hyperlinks (que no son runs normales)
        for hyper in p._element.findall(f'{ns_w}hyperlink'):
            p._element.remove(hyper)
    # Escribir en el primer párrafo
    para = celda.paragraphs[0]
    if para.runs:
        para.runs[0].text = texto
        if negrita:
            para.runs[0].bold = True
    else:
        run = para.add_run(texto)
        if negrita:
            run.bold = True


def escribir_firma(celda, personas):
    """
    Escribe firmas en una celda reutilizando la estructura del template.
    personas = [("TITULO", [("NOMBRE1", "CC1"), ("NOMBRE2", "CC2"), ...])]
    Preserva negrita, espaciado y formato original.
    """
    ns_w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    tc = celda._tc

    # Eliminar todos los párrafos existentes
    for p_el in tc.findall(f'{ns_w}p'):
        tc.remove(p_el)

    def agregar_parrafo(texto, negrita=False, sin_espacio=False):
        """Agrega un párrafo con un run al final de la celda."""
        from lxml import etree
        p_el = etree.SubElement(tc, f'{ns_w}p')
        # Eliminar espacio entre parrafos para que queden juntos
        if sin_espacio:
            ppr = etree.SubElement(p_el, f'{ns_w}pPr')
            spacing = etree.SubElement(ppr, f'{ns_w}spacing')
            spacing.set(f'{ns_w}after', '0')
            spacing.set(f'{ns_w}before', '0')
            spacing.set(f'{ns_w}line', '240')
            spacing.set(f'{ns_w}lineRule', 'auto')
        if texto:
            r_el = etree.SubElement(p_el, f'{ns_w}r')
            rpr = etree.SubElement(r_el, f'{ns_w}rPr')
            if negrita:
                etree.SubElement(rpr, f'{ns_w}b')
            t_el = etree.SubElement(r_el, f'{ns_w}t')
            t_el.text = texto
            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        return p_el

    titulo, lista_firmas = personas
    # Título (negrita)
    agregar_parrafo(titulo, negrita=True)

    for nombre, cc in lista_firmas:
        # Líneas vacías de separación antes de la firma
        agregar_parrafo("")
        agregar_parrafo("")
        # Línea de firma, nombre y CC juntos sin espacio vertical
        agregar_parrafo("__________________________________ ", negrita=True, sin_espacio=True)
        agregar_parrafo(nombre, negrita=True, sin_espacio=True)
        agregar_parrafo(f"C.C. No. {cc}", negrita=True, sin_espacio=True)


def generar_contrato_desde_formulario(datos_enriquecidos: dict, ruta_template: Path) -> Path:
    """
    Copia el template del contrato, escribe directamente en las tablas
    y reemplaza texto en cláusulas/pagarés con los datos del formulario.
    """
    todos_deudores = datos_enriquecidos["deudores"]
    deudor = todos_deudores[0] if todos_deudores else {}
    deudores_extra = todos_deudores[1:] if len(todos_deudores) > 1 else []
    codeudores = datos_enriquecidos.get("codeudores", [])
    acreedores = datos_enriquecidos["acreedores"]
    acr1 = acreedores[0] if len(acreedores) > 0 else {}
    acr2 = acreedores[1] if len(acreedores) > 1 else {}
    inm = datos_enriquecidos["inmueble"]
    prest = datos_enriquecidos["prestamo"]
    fecha = datos_enriquecidos["fecha_firma"]

    # Generar archivo
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    nombre_deudor = deudor["nombre_completo_mayuscula"].replace(" ", "_")[:30]
    nombre_archivo = f"Contrato_{nombre_deudor}_{ts}.docx"
    ruta_output = OUTPUT_DIR / nombre_archivo

    shutil.copy2(str(ruta_template), str(ruta_output))
    doc = Document(str(ruta_output))

    # ──────────────────────────────────────────
    # TABLA 0: DATOS BASICOS (28 filas x 2 cols)
    # ──────────────────────────────────────────
    t0 = doc.tables[0]

    # Fila 1: Fecha
    escribir_celda(t0, 1, 1, fecha)
    # Fila 3: Monto TOTAL del credito
    escribir_celda(t0, 3, 1, prest["monto_total_texto"])
    # Fila 4: Tasa
    escribir_celda(t0, 4, 1, f"{prest['tasa_mensual']} mensual anticipado")
    # Fila 5: Plazo
    escribir_celda(t0, 5, 1, f"{prest['plazo_texto']} contados a partir de la Fecha de Desembolso del Monto Inicial")
    # Fila 6: Matricula
    if inm.get("matricula_inmobiliaria"):
        escribir_celda(t0, 6, 1, inm["matricula_inmobiliaria"])
    # Fila 7: Direccion inmueble
    if inm.get("direccion_corta"):
        escribir_celda(t0, 7, 1, inm["direccion_corta"])

    # Fila 9-14: Deudor principal
    escribir_celda(t0, 9, 1, deudor["nombre_completo"], negrita=True)
    escribir_celda(t0, 10, 1, f"C.C. No. {deudor['cc']}")
    escribir_celda(t0, 11, 1, deudor["estado_civil"])
    escribir_celda(t0, 12, 1, deudor["direccion"])
    escribir_celda(t0, 13, 1, deudor["email"])
    escribir_celda(t0, 14, 1, deudor["telefono"])

    # ── Deudores adicionales: insertar filas después del principal (fila 14) ──
    fila_insercion = 14
    for ide, dex in enumerate(deudores_extra):
        encabezado = f"INFORMACIÓN DEL DEUDOR {ide + 2}"
        insertar_fila_tabla(t0, fila_insercion, encabezado, encabezado)
        fila_insercion += 1
        insertar_fila_tabla(t0, fila_insercion, "Nombre", dex["nombre_completo"], negrita_col1=True)
        fila_insercion += 1
        campos_dex = [
            ("Documento de identidad", f"C.C. No. {dex['cc']}"),
            ("Estado civil", dex.get("estado_civil", "")),
            ("Dirección de notificación", dex.get("direccion", "")),
            ("Correo electrónico", dex.get("email", "")),
            ("Número de celular", dex.get("telefono", "")),
        ]
        for etiqueta, valor in campos_dex:
            insertar_fila_tabla(t0, fila_insercion, etiqueta, valor)
            fila_insercion += 1

    # ── Codeudores: insertar filas después de todos los deudores ──
    # Cada codeudor necesita 7 filas: encabezado, nombre, CC, civil, dir, email, tel
    for ic, cod in enumerate(codeudores):
        encabezado = f"INFORMACIÓN DEL CODEUDOR {ic + 1}"
        insertar_fila_tabla(t0, fila_insercion, encabezado, encabezado)
        fila_insercion += 1
        insertar_fila_tabla(t0, fila_insercion, "Nombre", cod["nombre_completo"], negrita_col1=True)
        fila_insercion += 1
        campos_cod = [
            ("Documento de identidad", f"C.C. No. {cod['cc']}"),
            ("Estado civil", cod.get("estado_civil", "")),
            ("Dirección de notificación", cod.get("direccion", "")),
            ("Correo electrónico", cod.get("email", "")),
            ("Número de celular", cod.get("telefono", "")),
        ]
        for etiqueta, valor in campos_cod:
            insertar_fila_tabla(t0, fila_insercion, etiqueta, valor)
            fila_insercion += 1

    # Recalcular offset: las filas de acreedores se desplazaron
    filas_deudores_extra = len(deudores_extra) * 7
    filas_codeudores = len(codeudores) * 7
    offset = filas_deudores_extra + filas_codeudores

    # Filas de acreedores (originalmente 15-27, ahora desplazadas por offset)
    fila_acr1 = 16 + offset
    fila_acr2 = 22 + offset

    if acr1:
        escribir_celda(t0, fila_acr1, 1, acr1["nombre_completo"], negrita=True)
        escribir_celda(t0, fila_acr1 + 1, 1, f"C.C. No. {acr1['cc']}")
        escribir_celda(t0, fila_acr1 + 2, 1, acr1.get("estado_civil", ""))
        escribir_celda(t0, fila_acr1 + 3, 1, acr1.get("direccion", ""))
        escribir_celda(t0, fila_acr1 + 4, 1, acr1.get("email", ""))
        escribir_celda(t0, fila_acr1 + 5, 1, acr1.get("telefono", ""))
    if acr2:
        escribir_celda(t0, fila_acr2, 1, acr2["nombre_completo"], negrita=True)
        escribir_celda(t0, fila_acr2 + 1, 1, f"C.C. No. {acr2['cc']}")
        escribir_celda(t0, fila_acr2 + 2, 1, acr2.get("estado_civil", ""))
        escribir_celda(t0, fila_acr2 + 3, 1, acr2.get("direccion", ""))
        escribir_celda(t0, fila_acr2 + 4, 1, acr2.get("email", ""))
        escribir_celda(t0, fila_acr2 + 5, 1, acr2.get("telefono", ""))

    # ──────────────────────────────────────────
    # TABLA 1: DESCRIPCION DEL CREDITO (7 filas x 3 cols)
    # ──────────────────────────────────────────
    t1 = doc.tables[1]
    escribir_celda(t1, 0, 2, prest["monto_total_texto"])
    escribir_celda(t1, 2, 2, f"{prest['tasa_mensual']} mensual anticipado")
    escribir_celda(t1, 5, 2, str(prest.get("plazo_meses", "")))
    escribir_celda(t1, 6, 2, prest["cuota_mensual_total_texto"])

    # ──────────────────────────────────────────
    # TABLA 2: CUENTAS BANCARIAS (3 filas x 2 cols)
    # ──────────────────────────────────────────
    t2 = doc.tables[2]
    if acr1:
        escribir_celda(t2, 1, 0, acr1["nombre_completo"])
        escribir_celda(t2, 1, 1, acr1.get("cuenta_bancaria", ""))
    if acr2:
        escribir_celda(t2, 2, 0, acr2["nombre_completo"])
        escribir_celda(t2, 2, 1, acr2.get("cuenta_bancaria", ""))

    # ──────────────────────────────────────────
    # TABLA 3: FIRMAS CONTRATO (1 fila x 2 cols)
    # ──────────────────────────────────────────
    t3 = doc.tables[3]
    # Celda izquierda: Todos los deudores + Codeudores
    firmas_deudor = [(d["nombre_completo_mayuscula"], d["cc"]) for d in todos_deudores]
    for cod in codeudores:
        firmas_deudor.append((f"{cod['nombre_completo_mayuscula']}\n(CODEUDOR)", cod["cc"]))
    escribir_firma(t3.rows[0].cells[0], ("DEUDOR(ES)", firmas_deudor))
    # Celda derecha: Acreedores
    firmas_acr = [(a["nombre_completo"], a["cc"]) for a in acreedores]
    escribir_firma(t3.rows[0].cells[1], ("ACREEDORES", firmas_acr))

    # ──────────────────────────────────────────
    # TABLAS 4-7, 9: FIRMAS PAGARES Y CARTAS
    # ──────────────────────────────────────────
    for ti in [4, 5, 6, 7, 9]:
        if ti < len(doc.tables):
            tabla = doc.tables[ti]
            firmas_pag = [(d["nombre_completo"], d["cc"]) for d in todos_deudores]
            for cod in codeudores:
                firmas_pag.append((f"{cod['nombre_completo']}\n(CODEUDOR)", cod["cc"]))
            escribir_firma(tabla.rows[0].cells[0], ("DEUDOR(ES)", firmas_pag))

    # ──────────────────────────────────────────
    # TABLA 8: ANEXO 3 - INSTRUCCIONES DE ENTREGA (9 filas x 2 cols)
    # ──────────────────────────────────────────
    t8 = doc.tables[8]
    escribir_celda(t8, 1, 1, formato_pesos(prest["monto_total"]))
    escribir_celda(t8, 2, 1, formato_pesos(prest["monto_inicial_credito"]))
    escribir_celda(t8, 3, 1, formato_pesos(prest["comision_aluri_total"]))
    # Fila 6: Cuota mensual primer mes (pago anticipado)
    escribir_celda(t8, 6, 1, formato_pesos(prest["cuota_mensual_total"]))

    # ──────────────────────────────────────────
    # TABLA ANIDADA: PARTICIPACION DE ACREEDORES
    # (tabla nested dentro del cuerpo, entre PARAGRAFO TERCERO y CUARTO)
    # Encabezado: NOMBRE DEL ACREEDORES | % PARTICIPACIÓN | APORTE EN $COP
    # ──────────────────────────────────────────
    from docx.table import Table as DocxTable
    ns_w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    all_tbls = doc.element.body.findall(f'.//{ns_w}tbl')
    for tbl_el in all_tbls:
        # Buscar la tabla que tiene "NOMBRE DEL ACREEDOR" en la primera fila
        first_row = tbl_el.findall(f'{ns_w}tr')
        if not first_row:
            continue
        first_text = ''.join((t.text or '') for t in first_row[0].findall(f'.//{ns_w}t'))
        if 'NOMBRE DEL ACREEDOR' not in first_text.upper():
            continue

        tabla_part = DocxTable(tbl_el, doc)
        # Limpiar filas existentes (excepto encabezado)
        while len(tabla_part.rows) > 1:
            tbl_el.remove(tabla_part.rows[-1]._tr)

        # Agregar fila por cada acreedor
        for a in acreedores:
            pct = a.get("participacion_porcentaje", "")
            if pct and not pct.endswith("%"):
                pct = pct + "%"
            monto_aporte = formato_pesos(a.get("participacion_monto", 0))

            # Copiar estructura de la primera fila de datos (row 1 del template)
            from copy import deepcopy
            nueva_tr = deepcopy(first_row[1]) if len(first_row) > 1 else deepcopy(first_row[0])
            tcs = nueva_tr.findall(f'{ns_w}tc')
            valores = [a["nombre_completo"], pct, monto_aporte]
            for ci, tc in enumerate(tcs):
                if ci < len(valores):
                    for t in tc.findall(f'.//{ns_w}t'):
                        t.text = ""
                    ts = tc.findall(f'.//{ns_w}t')
                    if ts:
                        ts[0].text = valores[ci]
                    else:
                        # Crear un run con texto
                        p_el = tc.find(f'{ns_w}p')
                        if p_el is not None:
                            r_el = p_el.makeelement(f'{ns_w}r', {})
                            t_el = r_el.makeelement(f'{ns_w}t', {})
                            t_el.text = valores[ci]
                            r_el.append(t_el)
                            p_el.append(r_el)
            tbl_el.append(nueva_tr)

        # REGLA FIJA: tabla de participacion nunca se divide entre paginas
        # Aplicar cantSplit a cada fila para que Word las mantenga juntas
        for tr in tbl_el.findall(f'{ns_w}tr'):
            trPr = tr.find(f'{ns_w}trPr')
            if trPr is None:
                trPr = tr.makeelement(f'{ns_w}trPr', {})
                tr.insert(0, trPr)
            if trPr.find(f'{ns_w}cantSplit') is None:
                trPr.append(trPr.makeelement(f'{ns_w}cantSplit', {}))

        break

    # ──────────────────────────────────────────
    # TEXTO DE CLAUSULAS Y PAGARES (buscar y reemplazar en parrafos)
    # Usamos los nombres/datos del TEMPLATE para reemplazar
    # ──────────────────────────────────────────
    # Nombres template originales (para buscar en texto de clausulas)
    TPL_DEUDOR_NOMBRES = [
        "PATRICIA GONZÁLEZ VELÁSQUEZ",
        "PATRICIA GONZALEZ VELASQUEZ",
    ]
    TPL_ACR1_NOMBRES = [
        "HERNÁN JESÚS MONDOL CABARCAS",
        "HERNAN JESUS MONDOL CABARCAS",
    ]
    TPL_ACR2_NOMBRES = [
        "HAWER ALBERTO HERRERA RODRÍGUEZ",
        "HAWER ALBERTO HERRERA RODRIGUEZ",
    ]

    reemplazos = []

    # Deudor en texto
    for tpl_name in TPL_DEUDOR_NOMBRES:
        reemplazos.append((tpl_name, deudor["nombre_completo"]))
    reemplazos.append(("C.C. No. 52.202.940", f"C.C. No. {deudor['cc']}"))
    reemplazos.append(("52.202.940", deudor["cc"]))
    reemplazos.append(("mayor de edad y vecina del municipio de Cota",
                        f"mayor de edad y vecino/a del municipio de {deudor.get('municipio', '')}"))

    # Acreedor 1 en texto (pagares, clausulas)
    if acr1:
        for tpl_name in TPL_ACR1_NOMBRES:
            reemplazos.append((tpl_name, acr1["nombre_completo"]))
        reemplazos.append(("C.C. No. 73.160.553", f"C.C. No. {acr1['cc']}"))
        reemplazos.append(("73.160.553", acr1["cc"]))
        reemplazos.append(("cédula de ciudadanía número 73.160.553",
                            f"cédula de ciudadanía número {acr1['cc']}"))

    # Acreedor 2 en texto (pagares, clausulas)
    if acr2:
        for tpl_name in TPL_ACR2_NOMBRES:
            reemplazos.append((tpl_name, acr2["nombre_completo"]))
        reemplazos.append(("C.C. No. 93.365.845", f"C.C. No. {acr2['cc']}"))
        reemplazos.append(("93.365.845", acr2["cc"]))
        reemplazos.append(("cédula de ciudadanía número 93.365.845",
                            f"cédula de ciudadanía número {acr2['cc']}"))

    # Montos en texto (el contrato maneja montos TOTALES del credito)
    mt = prest["monto_total"]
    mi = prest["monto_inicial_credito"]
    mr = prest["monto_restante"]

    # Monto total del credito (dos formatos en el template: "MONEDA CORRIENTE" y "M/CTE")
    reemplazos.append(("NOVENTA MILLONES DE PESOS MONEDA CORRIENTE (COP$90.000.000)",
                        prest["monto_total_texto"]))
    # Anexo 3 usa formato "M/CTE"
    monto_texto_mcte = f"{numero_a_texto(mt).upper()} DE PESOS M/CTE (COP{formato_pesos(mt)})"
    reemplazos.append(("NOVENTA MILLONES DE PESOS M/CTE (COP$90.000.000)", monto_texto_mcte))
    reemplazos.append(("$90.000.000", formato_pesos(mt)))
    # 35% Monto Inicial del Credito
    reemplazos.append(("TREINTA Y UN MILLONES QUINIENTOS MIL PESOS MONEDA CORRIENTE (COP$31.500.000)",
                        monto_a_texto_legal(mi)))
    reemplazos.append(("$31.500.000", formato_pesos(mi)))
    # 65% Monto Restante del Credito
    reemplazos.append(("CINCUENTA Y OCHO MILLONES QUINIENTOS MIL PESOS MONEDA CORRIENTE (COP$58.500.000)",
                        monto_a_texto_legal(mr)))
    # Cuota mensual total
    reemplazos.append(("Un millón seiscientos veinte mil pesos moneda corriente (COP$1.620.000)",
                        prest["cuota_mensual_total_texto"]))
    # Cuota anticipada primer mes (Anexo 3, parrafo 402)
    reemplazos.append(("UN MILLÓN SETECIENTOS DIEZ MIL PESOS MONEDA CORRIENTE (COP$1.710.000)",
                        monto_a_texto_legal(prest["cuota_mensual_total"])))
    reemplazos.append(("$1.710.000", formato_pesos(prest["cuota_mensual_total"])))

    # Fechas en texto
    reemplazos.append(("Ocho (8) de abril de dos mil veintiséis (2026)", fecha))
    reemplazos.append(("ocho (8) de abril del año dos mil veintiséis (2026)", fecha.lower()))
    reemplazos.append(("ocho (8) días del mes de abril del año dos mil veintiséis (2026)", fecha.lower()))
    reemplazos.append(("ocho (8) del mes de abril del año dos mil veintiséis (2026)", fecha.lower()))
    reemplazos.append(("el día ocho (8) de abril del año dos mil veintiséis (2026)", f"el dia {fecha.lower()}"))

    # Participacion de acreedores (PARAGRAFO CUARTO)
    # Texto original: "cada uno cuenta con una participación equivalente al cincuenta por ciento (50%) del crédito"
    # Generar texto dinamico segun las participaciones reales
    if len(acreedores) >= 2:
        partes_texto = []
        for a in acreedores:
            pct = a.get("participacion_porcentaje", "").replace("%", "").strip()
            monto = a.get("participacion_monto", 0)
            if pct and monto:
                partes_texto.append(
                    f"{a['nombre_completo']} con una participación de {formato_pesos(monto)} "
                    f"equivalente al {pct}% del crédito"
                )
        if partes_texto:
            texto_participacion = ", y ".join(partes_texto)
            reemplazos.append((
                "cada uno cuenta con una participación equivalente al cincuenta por ciento (50%) del crédito",
                texto_participacion
            ))

    # Tasa y plazo en texto
    reemplazos.append(("1.80% mensual anticipado", f"{prest['tasa_mensual']} mensual anticipado"))
    reemplazos.append(("Sesenta (60) meses", prest["plazo_texto"]))

    # Inmueble en texto
    if inm.get("matricula_inmobiliaria"):
        reemplazos.append(("50S-604333", inm["matricula_inmobiliaria"]))
    if inm.get("direccion_corta"):
        reemplazos.append(("Calle 19 Sur No. 69A - 48, Bogotá D.C.", inm["direccion_corta"]))
    if inm.get("chip"):
        reemplazos.append(("AAA0041FAAF", inm["chip"]))

    # Aplicar reemplazos en todo el documento (párrafos + tablas)
    for buscar, reemplazar in reemplazos:
        if buscar and reemplazar and buscar != reemplazar:
            reemplazar_en_documento(doc, buscar, reemplazar)

    # ──────────────────────────────────────────────────────────
    # ELIMINAR PAGINAS VACIAS: quitar bloques de parrafos vacios
    # consecutivos (>= 5) que generan paginas en blanco.
    # Se reemplazan por un salto de pagina limpio.
    # ──────────────────────────────────────────────────────────
    from lxml import etree
    ns_w2 = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    body = doc.element.body

    # Eliminar bloques de 5+ parrafos vacios consecutivos
    # (se repite hasta que no queden mas)
    while True:
        consecutivos = 0
        inicio_bloque = 0
        encontrado = False
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == "":
                if consecutivos == 0:
                    inicio_bloque = i
                consecutivos += 1
            else:
                if consecutivos >= 5:
                    for j in range(inicio_bloque + consecutivos - 1, inicio_bloque - 1, -1):
                        body.remove(doc.paragraphs[j]._element)
                    encontrado = True
                    break
                consecutivos = 0
        if not encontrado and consecutivos >= 5:
            for j in range(inicio_bloque + consecutivos - 1, inicio_bloque - 1, -1):
                body.remove(doc.paragraphs[j]._element)
            encontrado = True
        if not encontrado:
            break

    # Limpiar TODOS los pageBreakBefore que vienen del template
    # (algunos son incorrectos, ej: parrafos de dacion en pago)
    for p in doc.paragraphs:
        pPr = p._element.find(f'{ns_w2}pPr')
        if pPr is not None:
            pb = pPr.find(f'{ns_w2}pageBreakBefore')
            if pb is not None:
                pPr.remove(pb)

    # Agregar saltos de pagina limpios donde corresponde
    for p in doc.paragraphs:
        texto = p.text.strip()
        necesita_salto = False
        # CLAUSULAS siempre empieza en pagina 2
        if texto in ("CLÁUSULAS", "CL\u00c1USULAS", "CLAUSULAS"):
            necesita_salto = True
        # Pagares
        elif texto.startswith("PAGAR") and "No." in texto:
            necesita_salto = True
        # Anexo No. 3 titulo (sin ":" para no confundir con la referencia)
        elif texto in ("Anexo No. 3", "Anexo No. 3 "):
            necesita_salto = True
        if necesita_salto:
            pPr = p._element.find(f'{ns_w2}pPr')
            if pPr is None:
                pPr = etree.SubElement(p._element, f'{ns_w2}pPr')
                p._element.insert(0, pPr)
            etree.SubElement(pPr, f'{ns_w2}pageBreakBefore')

    # ──────────────────────────────────────────────────────────
    # FORMATO DINAMICO: EVITAR TITULOS CORTADOS DE SU CONTENIDO
    # Aplica keepNext + keepLines a todos los titulos/encabezados
    # para que nunca queden separados del texto que les sigue,
    # sin importar la cantidad de personas en el contrato.
    # ──────────────────────────────────────────────────────────
    import re as re_mod
    from lxml import etree

    ns_w2 = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    body = doc.element.body

    # Patron que detecta titulos de clausulas, paragrafos y secciones
    patron_titulo = re_mod.compile(
        r"^("
        r"PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|"
        r"S[EÉ]PTIMA|OCTAVA|NOVENA|"
        r"D[EÉ]CIMA|VIG[EÉ]SIMA|TRIG[EÉ]SIMA|"
        r"PAR[AÁ]GRAFO|"
        r"CL[AÁ]USULAS|"
        r"Anexo No\.|"
        r"PAGAR[EÉ]|"
        r"CARTA DE INSTRUCCIONES|"
        r"INSTRUCCIONES DE ENTREGA|"
        r"INSTRUCCIONES$|"
        r"REFERENCIA:"
        r")",
        re_mod.IGNORECASE
    )

    def aplicar_keep(parrafo, keep_next=True, keep_lines=True):
        """Aplica keepNext y keepLines a un parrafo."""
        pPr = parrafo._element.find(f'{ns_w2}pPr')
        if pPr is None:
            pPr = etree.SubElement(parrafo._element, f'{ns_w2}pPr')
            parrafo._element.insert(0, pPr)
        if keep_next and pPr.find(f'{ns_w2}keepNext') is None:
            etree.SubElement(pPr, f'{ns_w2}keepNext')
        if keep_lines and pPr.find(f'{ns_w2}keepLines') is None:
            etree.SubElement(pPr, f'{ns_w2}keepLines')

    def aplicar_page_break(parrafo):
        """Aplica pageBreakBefore a un parrafo."""
        pPr = parrafo._element.find(f'{ns_w2}pPr')
        if pPr is None:
            pPr = etree.SubElement(parrafo._element, f'{ns_w2}pPr')
            parrafo._element.insert(0, pPr)
        if pPr.find(f'{ns_w2}pageBreakBefore') is None:
            etree.SubElement(pPr, f'{ns_w2}pageBreakBefore')

    # Aplicar keepNext a todos los titulos para que no se separen de su texto
    for p in doc.paragraphs:
        texto = p.text.strip()
        if not texto:
            continue
        if patron_titulo.match(texto):
            aplicar_keep(p, keep_next=True, keep_lines=True)

    # ── Saltos de pagina fijos para secciones principales ──

    # Salto antes de la lista de anexos ("Anexo No. 1: Escritura...")
    for p in doc.paragraphs:
        if "Anexo No. 1" in p.text and "Escritura" in p.text:
            aplicar_page_break(p)
            break

    doc.save(str(ruta_output))
    return ruta_output


# ═══════════════════════════════════════════════
# RUTAS FLASK
# ═══════════════════════════════════════════════

@app.route("/")
def index():
    return render_template("checklist_form.html")


@app.route("/api/cargar-checklist", methods=["POST"])
def cargar_checklist():
    """
    Carga un Check_List.docx (subido por el usuario o el default) y extrae los datos
    para rellenar el formulario web.
    """
    # Si viene un archivo subido, usarlo; si no, usar el default
    if "archivo" in request.files and request.files["archivo"].filename:
        archivo = request.files["archivo"]
        tmp_path = DATA_DIR / f"_tmp_upload_{datetime.now().strftime('%H%M%S')}.docx"
        archivo.save(str(tmp_path))
        ruta = tmp_path
    elif CHECKLIST_PATH.exists():
        ruta = CHECKLIST_PATH
    else:
        return jsonify({"ok": False, "error": "No se encontro el archivo Check_List.docx"}), 404

    try:
        datos = parsear_checklist_docx(ruta)
        return jsonify({"ok": True, "datos": datos})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        # Limpiar archivo temporal si existio
        if "tmp_path" in dir() and tmp_path.exists() and tmp_path != CHECKLIST_PATH:
            try:
                tmp_path.unlink()
            except OSError:
                pass


def parsear_checklist_docx(ruta: Path) -> dict:
    """
    Lee un Check_List.docx y extrae todos los campos en un diccionario
    compatible con la estructura del formulario web.
    """
    doc = Document(str(ruta))

    # Extraer todas las lineas con texto
    lineas = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lineas.append(t)

    texto_completo = "\n".join(lineas)

    def buscar(patron, texto=texto_completo, grupo=1):
        """Busca un patron regex y retorna el grupo indicado o ''."""
        m = re.search(patron, texto, re.IGNORECASE)
        return m.group(grupo).strip() if m else ""

    def buscar_campo(etiqueta, texto=texto_completo):
        """Busca 'Etiqueta: valor' en el texto."""
        m = re.search(rf"{etiqueta}\s*:\s*(.+)", texto, re.IGNORECASE)
        return m.group(1).strip() if m else ""

    # ── Tipo de contrato ──
    tipo = buscar_campo("TIPO DE CONTRATO") or "Hipoteca"

    # ──────────────────────────────────────────────
    # PARSEO SECUENCIAL DE BLOQUES DE PERSONAS
    # Recorremos linea por linea desde el inicio hasta "Acreedor 1"
    # Cada bloque empieza con un encabezado (DEUDOR:, CODEUDOR N:,
    # "SE LE COMPRA...", o cualquier linea con link/titulo que no sea
    # un campo Nombre/CC/etc.) y contiene los campos de esa persona.
    # ──────────────────────────────────────────────

    def extraer_persona_de_bloque(bloque_texto):
        """Extrae los datos de una persona desde un bloque de texto."""
        nombre = buscar_campo("Nombre", bloque_texto)
        if not nombre or re.match(r"^(CC|Cedula|Direcci|Correo|Tel|Estado|Participaci)", nombre, re.IGNORECASE):
            return None  # Bloque vacio

        cc_raw = buscar(r"CC\.?\s*(?:del\s*Deudor)?\s*:\s*([^\n]+)", bloque_texto)
        cc, cc_exp = separar_cc_y_expedicion(limpiar_campo_checklist(cc_raw))
        direccion = limpiar_campo_checklist(buscar(r"Direcci.n\s*(?:de\s*)?notificaci.n\s*(?:Deudor)?\s*:\s*([^\n]+)", bloque_texto))
        email = limpiar_campo_checklist(buscar(r"Correo\s*(?:electr.nico)?\s*(?:Deudor)?\s*:\s*([^\n]+)", bloque_texto))
        telefono = limpiar_campo_checklist(buscar(r"Tel.fono\s*(?:Deudor)?\s*:\s*([^\n]+)", bloque_texto))
        civil = limpiar_campo_checklist(buscar(r"Estado\s*civil\s*(?:Deudor)?\s*:\s*([^\n]+)", bloque_texto))
        part_monto_raw = buscar(r"Participaci.n\s*\$+\s*:\s*([^\n]+)", bloque_texto)
        part_pct = buscar(r"Participaci.n\s*%\s*:\s*([^\n]+)", bloque_texto)

        return {
            "nombre": nombre,
            "cc": cc,
            "cc_expedicion": cc_exp,
            "direccion": direccion,
            "email": email,
            "telefono": telefono,
            "estado_civil": civil,
            "participacion_monto": formatear_monto_display(limpiar_campo_checklist(part_monto_raw).replace("$", "").strip()),
            "participacion_porcentaje": limpiar_campo_checklist(part_pct),
        }

    # Cortar el texto entre el inicio y "Acreedor 1" para obtener
    # todos los bloques de deudor/codeudores
    m_corte = re.search(r"Acreedor\s+1", texto_completo, re.IGNORECASE)
    zona_personas = texto_completo[:m_corte.start()] if m_corte else texto_completo

    # Dividir en bloques: cada bloque empieza con un encabezado reconocible
    # Encabezados: "DEUDOR:", "CODEUDOR N:", "SE LE COMPRA...", o un link/titulo
    # que precede un bloque de campos Nombre/CC/etc.
    patron_encabezado = re.compile(
        r"^(DEUDOR\s*:|CODEUDOR\s*\d*\s*:|SE LE COMPRA|https?://)",
        re.IGNORECASE | re.MULTILINE
    )

    encabezados = list(patron_encabezado.finditer(zona_personas))

    deudores_lista = []
    codeudores = []

    for idx_enc, match_enc in enumerate(encabezados):
        inicio = match_enc.start()
        fin = encabezados[idx_enc + 1].start() if idx_enc + 1 < len(encabezados) else len(zona_personas)
        bloque = zona_personas[inicio:fin]

        # Determinar tipo segun el encabezado
        encab = match_enc.group(0).upper().strip()
        es_codeudor = encab.startswith("CODEUDOR")
        # DEUDOR, SE LE COMPRA, https (sin CODEUDOR) = deudor
        # CODEUDOR = codeudor

        persona = extraer_persona_de_bloque(bloque)
        if persona is None:
            continue

        if es_codeudor:
            codeudores.append(persona)
        else:
            deudores_lista.append(persona)

    # Si no se encontro ningun deudor, crear uno vacio
    if not deudores_lista:
        deudores_lista = [{
            "nombre": "", "cc": "", "cc_expedicion": "", "direccion": "",
            "email": "", "telefono": "", "estado_civil": "",
            "participacion_monto": "", "participacion_porcentaje": "",
        }]

    deudor_datos = deudores_lista[0]

    # ── Acreedores ──
    acreedores = []
    for i in range(1, 5):
        patron_bloque = rf"Acreedor\s+{i}\s*:?\s*\n?(.*?)(?=Acreedor\s+{i+1}|Inmueble|Condiciones|$)"
        m_acr = re.search(patron_bloque, texto_completo, re.DOTALL | re.IGNORECASE)
        if not m_acr:
            continue
        bloque = m_acr.group(0)

        nombre = buscar_campo("Nombre", bloque)
        if not nombre:
            continue  # Acreedor vacio
        # Descartar si el "nombre" es en realidad otra etiqueta (acreedor vacio)
        if re.match(r"^(Cedula|Direccion|Correo|Telefono|Estado|Participacion)", nombre, re.IGNORECASE):
            continue

        cc_raw = buscar(r"[Cc].dula\s*:\s*(.+)", bloque)
        cc, cc_exp = separar_cc_y_expedicion(cc_raw)

        direccion = buscar(r"[Dd]irecci.n\s*notificaci?o?n\s*:\s*(.+)", bloque)
        email = buscar(r"[Cc]orreo\s*:\s*(.+)", bloque)
        telefono = buscar(r"[Tt]el.fono\s*:\s*(.+)", bloque)
        civil = buscar(r"[Ee]stado [Cc]ivil\s*:\s*(.+)", bloque).strip()
        part_monto = buscar(r"[Pp]articipaci.n\s*\$+\s*:\s*(.+)", bloque).replace("$", "").replace(" ", "").strip()
        part_pct = buscar(r"[Pp]articipaci.n\s*%\s*:\s*(.+)", bloque)

        acreedores.append({
            "nombre": nombre,
            "cc": cc,
            "cc_expedicion": cc_exp,
            "direccion": direccion,
            "email": email,
            "telefono": telefono,
            "estado_civil": civil,
            "participacion_monto": formatear_monto_display(part_monto),
            "participacion_porcentaje": part_pct,
            "cuenta_bancaria": "",
        })

    # ── Inmueble ──
    bloque_inmueble = ""
    m_inm = re.search(r"Inmueble\s*:?\s*\n?(.*?)(?=Condiciones|$)", texto_completo, re.DOTALL | re.IGNORECASE)
    if m_inm:
        bloque_inmueble = m_inm.group(0)

    matricula = buscar(r"matr.cula inmobiliaria.*?:\s*(.+)", bloque_inmueble)
    matricula = matricula.rstrip(". ")
    cedula_catastral = buscar(r"[Cc].dula catastral.*?:\s*(.+)", bloque_inmueble)
    cedula_catastral = cedula_catastral.rstrip(". ")
    chip = buscar(r"CHIP\s*:\s*(.+)", bloque_inmueble)
    inmueble_dir = buscar(r"Direcci.n del [Ii]nmueble\s*:\s*(.+)", bloque_inmueble)
    inmueble_desc = buscar(r"Descripci.n del [Ii]nmueble\s*:\s*(.+)", bloque_inmueble)
    inmueble_linderos = buscar(r"Linderos\s*:\s*(.+)", bloque_inmueble)

    # Concatenar lineas de linderos (suelen estar en multiples parrafos)
    linderos_lineas = []
    capturando = False
    for linea in lineas:
        if re.match(r"Linderos\s*:", linea, re.IGNORECASE):
            capturando = True
            rest = re.sub(r"Linderos\s*:\s*", "", linea, flags=re.IGNORECASE).strip()
            if rest:
                linderos_lineas.append(rest)
            continue
        if capturando:
            if re.match(r"(C.digo CHIP|Condiciones|Monto)", linea, re.IGNORECASE):
                break
            if re.match(r"POR EL", linea, re.IGNORECASE) or (linderos_lineas and not re.match(r"\w+\s*:", linea)):
                linderos_lineas.append(linea)
            else:
                break
    if linderos_lineas:
        inmueble_linderos = " ".join(linderos_lineas)

    # ── Condiciones del prestamo ──
    bloque_prest = ""
    m_prest = re.search(r"Condiciones del pr.stamo\s*:?\s*\n?(.*?)(?=Observaci.n|$)", texto_completo, re.DOTALL | re.IGNORECASE)
    if m_prest:
        bloque_prest = m_prest.group(0)

    monto = buscar(r"Monto del pr.stamo\s*:\s*\$?\s*(.+)", bloque_prest).replace("$", "").replace(" ", "").strip()
    plazo = buscar(r"Plazo\s*\(?meses?\)?\s*:\s*(\d+)", bloque_prest)
    tasa = buscar(r"Tasa\s*\(?\s*mes\s*anticipado\s*\)?\s*:\s*(.+)", bloque_prest)
    cuota = buscar(r"Valor de la cuota mensual\s*:\s*\$?\s*(.+)", bloque_prest).replace("$", "").replace(" ", "").strip()
    forma_pago_raw = buscar(r"Forma de pago.*?:\s*(.+)", bloque_prest)
    comision = buscar(r"Comisi.n Aluri\s*:\s*\$?\s*(.+)", bloque_prest).replace("$", "").replace(" ", "").strip()

    # Normalizar forma de pago
    forma_pago = "Solo intereses"
    if forma_pago_raw:
        fp_lower = forma_pago_raw.lower()
        if "capital" in fp_lower and "inter" in fp_lower:
            forma_pago = "Interes y capital"
        elif "solo" in fp_lower and "inter" in fp_lower:
            forma_pago = "Solo intereses"

    # ── Observaciones ──
    observaciones = buscar(r"Observaci.n\s*:\s*(.+)", texto_completo)

    return {
        "tipo_contrato": tipo,
        "deudor": deudor_datos,
        "deudores": deudores_lista,
        "codeudores": codeudores,
        "acreedores": acreedores,
        "inmueble": {
            "matricula_inmobiliaria": matricula,
            "cedula_catastral": cedula_catastral,
            "chip": chip,
            "direccion": inmueble_dir,
            "descripcion": inmueble_desc,
            "linderos": inmueble_linderos,
        },
        "prestamo": {
            "monto": formatear_monto_display(monto),
            "plazo_meses": plazo,
            "tasa_mensual": tasa,
            "cuota_mensual": formatear_monto_display(cuota),
            "forma_pago": forma_pago,
            "comision_aluri": formatear_monto_display(comision),
            "observaciones": observaciones,
        },
    }


def limpiar_campo_checklist(valor: str) -> str:
    """
    Si el valor capturado por el regex es en realidad otra etiqueta
    del checklist (porque el campo estaba vacio), retorna ''.
    """
    if not valor:
        return ""
    v = valor.strip()
    # Si empieza con una etiqueta conocida, el campo estaba vacio
    if re.match(r"^(Participaci|Nombre|Cedula|Direcci|Correo|Tel.fono|Estado|CODEUDOR|Acreedor|Inmueble|Condiciones|Monto|SE LE|DEUDOR)", v, re.IGNORECASE):
        return ""
    return v


def separar_cc_y_expedicion(cc_raw: str) -> tuple:
    """
    Separa '52.202.940 de Bogota' en ('52.202.940', 'Bogota').
    Tambien maneja 'C.C. No. 1.026.550.415 Bogota D.C.' -> ('1.026.550.415', 'Bogota D.C.')
    Si solo hay numeros, retorna (numero, '').
    """
    if not cc_raw:
        return "", ""
    # Quitar prefijos como "C.C.", "C.C. No.", "CC.", "No."
    limpio = re.sub(r"^(C\.?C\.?\s*(No\.?)?\s*)", "", cc_raw.strip(), flags=re.IGNORECASE).strip()
    # Patron: numero + "de" + ciudad
    m = re.match(r"([\d.]+)\s+de\s+(.+)", limpio)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    # Patron: numero + espacio + ciudad (sin "de")
    solo_num = re.match(r"([\d.]+)\s*(.*)", limpio)
    if solo_num:
        cc = solo_num.group(1).strip()
        resto = solo_num.group(2).strip()
        return cc, resto
    return limpio, ""



def interpretar_monto(valor: str) -> int:
    """
    Interpreta un monto que puede venir como:
      - '180.000.000' o '180000000' -> 180000000
      - '93 MILLONES' o '93 millones' -> 93000000
      - '1.5 MILLONES' -> 1500000
      - '500 MIL' -> 500000
      - '$90.000.000' -> 90000000
    Retorna el valor como entero.
    """
    if not valor:
        return 0
    texto = valor.strip().upper().replace("$", "").replace(",", "").strip()

    # Patron: numero + "MILLONES" / "MILLON" / "MIL"
    m = re.match(r"([\d.]+)\s*(MILLONES|MILLON|MIL)", texto)
    if m:
        numero = float(m.group(1).replace(".", "", m.group(1).count(".") - 1) if m.group(1).count(".") > 1 else m.group(1))
        unidad = m.group(2)
        if unidad in ("MILLONES", "MILLON"):
            return int(numero * 1_000_000)
        elif unidad == "MIL":
            return int(numero * 1_000)

    # Patron numerico puro: quitar puntos de miles
    limpio = texto.replace(".", "").replace(" ", "").strip()
    try:
        return int(limpio)
    except ValueError:
        return 0


def formatear_monto_display(valor: str) -> str:
    """
    Interpreta el monto (soporta '93 MILLONES', '180.000.000', etc.)
    y lo formatea como '93.000.000' para mostrar en el formulario.
    """
    if not valor:
        return ""
    n = interpretar_monto(valor)
    if n > 0:
        return f"{n:,}".replace(",", ".")
    # Si no se pudo interpretar, devolver limpio
    limpio = valor.replace("$", "").replace(" ", "").strip()
    return limpio


@app.route("/api/generar-contrato", methods=["POST"])
def api_generar_contrato():
    """Genera el contrato .docx a partir de los datos del formulario y lo devuelve para descarga."""
    datos_form = request.get_json()
    if not datos_form:
        return jsonify({"ok": False, "error": "No se recibieron datos"}), 400

    if not TEMPLATE_PATH.exists():
        return jsonify({"ok": False, "error": "Template del contrato no encontrado"}), 500

    try:
        # Enriquecer datos (convertir numeros a texto, etc.)
        datos_enriquecidos = enriquecer_datos(datos_form)

        # Guardar checklist como JSON
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre = datos_form.get("deudor", {}).get("nombre", "sin_nombre").replace(" ", "_")[:30]
        json_path = DATA_DIR / f"checklist_{nombre}_{ts}.json"
        with open(str(json_path), "w", encoding="utf-8") as f:
            json.dump({"formulario": datos_form, "enriquecido": datos_enriquecidos}, f, ensure_ascii=False, indent=2)

        # Generar contrato
        ruta_contrato = generar_contrato_desde_formulario(datos_enriquecidos, TEMPLATE_PATH)

        return send_file(
            str(ruta_contrato),
            as_attachment=True,
            download_name=ruta_contrato.name,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/generar-contrato-pdf", methods=["POST"])
def api_generar_contrato_pdf():
    """Genera un PDF del formulario (Check List) con los datos ingresados."""
    datos_form = request.get_json()
    if not datos_form:
        return jsonify({"ok": False, "error": "No se recibieron datos"}), 400

    try:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre = datos_form.get("deudor", {}).get("nombre", "sin_nombre").replace(" ", "_")[:30]
        ruta_pdf = OUTPUT_DIR / f"Formulario_{nombre}_{ts}.pdf"
        OUTPUT_DIR.mkdir(exist_ok=True)

        generar_formulario_pdf(datos_form, ruta_pdf)

        return send_file(
            str(ruta_pdf),
            as_attachment=True,
            download_name=ruta_pdf.name,
            mimetype="application/pdf",
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"ok": False, "error": str(e)}), 500


def generar_formulario_pdf(datos: dict, ruta_salida: Path):
    """Genera un PDF con la informacion del formulario Check List."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    doc_pdf = SimpleDocTemplate(
        str(ruta_salida), pagesize=letter,
        leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm,
    )

    styles = getSampleStyleSheet()
    titulo_style = ParagraphStyle("Titulo", parent=styles["Heading1"], fontSize=14, alignment=1, spaceAfter=6)
    subtitulo_style = ParagraphStyle("Subtitulo", parent=styles["Heading2"], fontSize=11, textColor=colors.white, spaceAfter=2)
    normal_style = ParagraphStyle("Normal2", parent=styles["Normal"], fontSize=9, leading=11)
    small_style = ParagraphStyle("Small", parent=styles["Normal"], fontSize=8, leading=10, textColor=colors.grey)

    elements = []

    # Titulo
    elements.append(Paragraph("FORMULARIO DE SOLICITUD", titulo_style))
    elements.append(Paragraph("Check List - Informacion Requerida", small_style))
    elements.append(Spacer(1, 10))

    def seccion(titulo):
        t = Table([[Paragraph(titulo, subtitulo_style)]], colWidths=[17*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#1a3a5c")),
            ("TEXTCOLOR", (0, 0), (-1, -1), colors.white),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 4))

    def tabla_campos(campos):
        """campos = lista de (etiqueta, valor)"""
        data = []
        for etiq, val in campos:
            data.append([
                Paragraph(f"<b>{etiq}:</b>", normal_style),
                Paragraph(str(val or ""), normal_style),
            ])
        t = Table(data, colWidths=[6*cm, 11*cm])
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f4f6f9")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 8))

    # ── Tipo de contrato ──
    seccion("TIPO DE CONTRATO")
    tabla_campos([("Tipo de contrato", datos.get("tipo_contrato", ""))])

    # ── Deudores ──
    deudores = datos.get("deudores", [])
    if not deudores and datos.get("deudor"):
        deudores = [datos["deudor"]]

    for i, d in enumerate(deudores):
        label = "DEUDOR PRINCIPAL" if i == 0 else f"DEUDOR {i + 1}"
        seccion(f"INFORMACION DEL {label}")
        tabla_campos([
            ("Nombre", d.get("nombre", "")),
            ("No. Cedula", d.get("cc", "")),
            ("Expedida en", d.get("cc_expedicion", "")),
            ("Direccion", d.get("direccion", "")),
            ("Correo", d.get("email", "")),
            ("Telefono", d.get("telefono", "")),
            ("Estado civil", d.get("estado_civil", "")),
            ("Participacion $", d.get("participacion_monto", "")),
            ("Participacion %", d.get("participacion_porcentaje", "")),
        ])

    # ── Codeudores ──
    codeudores = datos.get("codeudores", [])
    if codeudores:
        for i, c in enumerate(codeudores):
            seccion(f"CODEUDOR {i + 1}")
            tabla_campos([
                ("Nombre", c.get("nombre", "")),
                ("No. Cedula", c.get("cc", "")),
                ("Expedida en", c.get("cc_expedicion", "")),
                ("Direccion", c.get("direccion", "")),
                ("Correo", c.get("email", "")),
                ("Telefono", c.get("telefono", "")),
                ("Estado civil", c.get("estado_civil", "")),
            ])

    # ── Acreedores ──
    for i, a in enumerate(datos.get("acreedores", [])):
        seccion(f"ACREEDOR {i + 1}")
        tabla_campos([
            ("Nombre", a.get("nombre", "")),
            ("No. Cedula", a.get("cc", "")),
            ("Expedida en", a.get("cc_expedicion", "")),
            ("Direccion", a.get("direccion", "")),
            ("Correo", a.get("email", "")),
            ("Telefono", a.get("telefono", "")),
            ("Estado civil", a.get("estado_civil", "")),
            ("Participacion $", a.get("participacion_monto", "")),
            ("Participacion %", a.get("participacion_porcentaje", "")),
            ("Cuenta bancaria", a.get("cuenta_bancaria", "")),
        ])

    # ── Inmueble ──
    inm = datos.get("inmueble", {})
    seccion("INFORMACION DEL INMUEBLE")
    tabla_campos([
        ("Matricula inmobiliaria", inm.get("matricula_inmobiliaria", "")),
        ("Cedula catastral", inm.get("cedula_catastral", "")),
        ("Codigo CHIP", inm.get("chip", "")),
        ("Direccion", inm.get("direccion", "")),
        ("Descripcion", inm.get("descripcion", "")),
        ("Linderos", inm.get("linderos", "")),
    ])

    # ── Prestamo ──
    p = datos.get("prestamo", {})
    seccion("CONDICIONES DEL PRESTAMO")
    tabla_campos([
        ("Monto del prestamo", p.get("monto", "")),
        ("Plazo (meses)", p.get("plazo_meses", "")),
        ("Tasa mensual", p.get("tasa_mensual", "")),
        ("Cuota mensual", p.get("cuota_mensual", "")),
        ("Forma de pago", p.get("forma_pago", "")),
        ("Comision Aluri", p.get("comision_aluri", "")),
        ("Observaciones", p.get("observaciones", "")),
    ])

    doc_pdf.build(elements)


@app.route("/api/listar-checklists", methods=["GET"])
def listar_checklists():
    """Lista todos los Check_Lists guardados."""
    archivos = sorted(DATA_DIR.glob("checklist_*.json"), reverse=True)
    resultado = []
    for a in archivos:
        with open(str(a), "r", encoding="utf-8") as f:
            datos = json.load(f)
        d = datos.get("formulario", datos)
        resultado.append({
            "archivo": a.name,
            "deudor": d.get("deudor", {}).get("nombre", ""),
            "monto": d.get("prestamo", {}).get("monto", ""),
            "fecha_creacion": a.stat().st_mtime,
        })
    return jsonify(resultado)


if __name__ == "__main__":
    print("\n  Formulario Check_List disponible en: http://localhost:5000\n")
    app.run(debug=True, port=5000)
