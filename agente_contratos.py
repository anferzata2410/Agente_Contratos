#!/usr/bin/env python3
"""
Agente de Contratos - Genera "Contrato Mutuo Solo Interés" a partir del Check_List

Uso:
    python agente_contratos.py
    python agente_contratos.py --checklist ruta/al/Check_List.docx
    python agente_contratos.py --checklist ruta/al/Check_List.docx --output ruta/salida.docx
"""

import re
import sys
import os
import json
import shutil
import argparse
import traceback
from pathlib import Path
from datetime import date

# Forzar UTF-8 en stdout para Windows
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
import anthropic

# Cargar .env si existe (sin depender de python-dotenv)
_env_path = Path(__file__).parent / ".env"
if _env_path.exists():
    for _line in _env_path.read_text(encoding="utf-8").splitlines():
        _line = _line.strip()
        if _line and not _line.startswith("#") and "=" in _line:
            _k, _v = _line.split("=", 1)
            os.environ.setdefault(_k.strip(), _v.strip())

# ─────────────────────────────────────────────
# Rutas por defecto
# ─────────────────────────────────────────────
BASE_DIR = Path(__file__).parent
DOCUMENTACION = BASE_DIR / "Documentacion"
DEFAULT_CHECKLIST = DOCUMENTACION / "Check_List.docx"
DEFAULT_TEMPLATE = DOCUMENTACION / "Contrato mutuo solo interés.docx"
DEFAULT_OUTPUT = BASE_DIR / "Contratos_Generados"


# ─────────────────────────────────────────────
# 1. EXTRACCIÓN DE DATOS CON CLAUDE API
# ─────────────────────────────────────────────

def leer_texto_checklist(ruta: Path) -> str:
    """Lee todo el texto del Check_List.docx y lo retorna como string."""
    doc = Document(str(ruta))
    lineas = []
    for para in doc.paragraphs:
        if para.text.strip():
            lineas.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() and cell.text.strip() not in lineas:
                    lineas.append(cell.text.strip())
    return "\n".join(lineas)


def extraer_datos_con_claude(texto_checklist: str) -> dict:
    """
    Usa la API de Claude para extraer datos estructurados del Check_List.
    Retorna un diccionario con toda la información necesaria para el contrato.
    """
    client = anthropic.Anthropic()

    prompt = f"""Analiza el siguiente formulario de solicitud de crédito hipotecario (Check_List) y extrae TODA la información en formato JSON estructurado exactamente como se indica.

TEXTO DEL CHECK_LIST:
{texto_checklist}

Extrae la información y responde ÚNICAMENTE con un JSON válido con esta estructura exacta:

{{
  "fecha_firma": "Fecha de hoy en texto: 'Ocho (8) de abril de dos mil veintiséis (2026)'",
  "deudores": [
    {{
      "nombre_completo": "PATRICIA GONZÁLEZ VELÁSQUEZ",
      "nombre_completo_mayuscula": "PATRICIA GONZALEZ VELASQUEZ",
      "cc": "52.202.940",
      "cc_expedicion": "Bogotá",
      "direccion": "Km 1.5 vía Cota Chia, casa 30, Cota",
      "email": "patricia.gonzalez.velasquez@hotmail.com",
      "telefono": "+57 3005821543",
      "estado_civil": "Casada con sociedad conyugal vigente",
      "municipio": "Cota"
    }}
  ],
  "acreedores": [
    {{
      "nombre_completo": "HERNÁN JESÚS MONDOL CABARCAS",
      "nombre_completo_mayuscula": "HERNAN JESUS MONDOL CABARCAS",
      "cc": "73.160.553",
      "cc_expedicion": "Cartagena",
      "direccion": "Cra 103b # 150C-30, casa 110, Agrupación Arboleda La Campiña, Bogotá D.C.",
      "email": "hernanmondol73@gmail.com",
      "telefono": "+57 3107668525",
      "estado_civil": "Soltero con unión marital de hecho",
      "participacion_porcentaje": "50%",
      "participacion_monto": 90000000,
      "participacion_texto": "NOVENTA MILLONES DE PESOS MONEDA CORRIENTE (COP$90.000.000)",
      "cuenta_bancaria": "Cuenta de ahorros número XXXXXXXXX de Bancolombia",
      "cuota_mensual_individual": 1620000,
      "cuota_mensual_texto": "Un millón seiscientos veinte mil pesos moneda corriente (COP$1.620.000)"
    }},
    {{
      "nombre_completo": "HAWER ALBERTO HERRERA RODRÍGUEZ",
      "nombre_completo_mayuscula": "HAWER ALBERTO HERRERA RODRIGUEZ",
      "cc": "93.365.845",
      "cc_expedicion": "Ibagué",
      "direccion": "Carrera 6 # 53-29, Of. 604, Ibagué",
      "email": "arqui.hawerh@gmail.com",
      "telefono": "+57 3157909502",
      "estado_civil": "Casado con sociedad conyugal vigente",
      "participacion_porcentaje": "50%",
      "participacion_monto": 90000000,
      "participacion_texto": "NOVENTA MILLONES DE PESOS MONEDA CORRIENTE (COP$90.000.000)",
      "cuenta_bancaria": "Cuenta de ahorros número XXXXXXXXX de Bancolombia",
      "cuota_mensual_individual": 1620000,
      "cuota_mensual_texto": "Un millón seiscientos veinte mil pesos moneda corriente (COP$1.620.000)"
    }}
  ],
  "inmueble": {{
    "matricula_inmobiliaria": "50S-604333",
    "oficina_registro": "Bogotá, Zona Sur",
    "cedula_catastral": "BS 23S 61 44 2",
    "chip": "AAA0041FAAF",
    "direccion_corta": "Calle 19 Sur No. 69A - 48, Bogotá D.C.",
    "descripcion_completa": "APARTAMENTO NÚMERO SESENTA Y DOS CUARENTA Y OCHO (62-48)...",
    "descripcion_larga": "texto completo de la descripción e linderos"
  }},
  "prestamo": {{
    "monto_total": 180000000,
    "monto_total_texto": "CIENTO OCHENTA MILLONES DE PESOS MONEDA CORRIENTE (COP$180.000.000)",
    "monto_inicial_credito": 63000000,
    "monto_inicial_texto": "SESENTA Y TRES MILLONES DE PESOS MONEDA CORRIENTE (COP$63.000.000)",
    "monto_inicial_porcentaje": "35%",
    "monto_restante": 117000000,
    "monto_restante_texto": "CIENTO DIECISIETE MILLONES DE PESOS MONEDA CORRIENTE (COP$117.000.000)",
    "plazo_meses": 60,
    "plazo_texto": "Sesenta (60) meses",
    "tasa_mensual": "1.80%",
    "tasa_texto": "1.80% mensual anticipado",
    "cuota_mensual_total": 3240000,
    "cuota_mensual_total_texto": "Tres millones doscientos cuarenta mil pesos moneda corriente (COP$3.240.000)",
    "forma_pago": "Solo intereses",
    "comision_aluri_total": 9000000,
    "comision_aluri_por_acreedor": 4500000,
    "servicios_aluri_texto": "$ 4.500.000"
  }},
  "observaciones": "texto de observaciones si hay"
}}

INSTRUCCIONES IMPORTANTES:
1. Extrae los datos exactamente como aparecen en el formulario
2. Para los montos del desembolso: el "Monto Inicial" (35%) es lo que se entrega con la boleta de registro. El "Monto Restante" (65%) se entrega después de verificar el certificado de libertad.
3. Para estado civil: adapta a "Casada/o con sociedad conyugal vigente", "Soltero/a con unión marital de hecho", "Soltero/a", etc.
4. Si no encuentras un dato, usa null
5. Responde SOLO con el JSON, sin texto adicional"""

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}]
    )

    respuesta = message.content[0].text.strip()

    # Limpiar markdown si viene envuelto en ```json
    if respuesta.startswith("```"):
        respuesta = re.sub(r"^```(?:json)?\n?", "", respuesta)
        respuesta = re.sub(r"\n?```$", "", respuesta)

    datos = json.loads(respuesta)
    return datos


# ─────────────────────────────────────────────
# 2. MANIPULACIÓN DEL DOCUMENTO WORD
# ─────────────────────────────────────────────

def reemplazar_texto_run(run, texto_nuevo: str):
    """Reemplaza el texto de un run preservando su formato."""
    run.text = texto_nuevo


def reemplazar_en_parrafo(parrafo, buscar: str, reemplazar: str) -> bool:
    """
    Reemplaza texto en un párrafo, reconstruyendo runs si es necesario.
    Retorna True si hizo algún reemplazo.
    """
    texto_completo = parrafo.text
    if buscar not in texto_completo:
        return False

    # Caso simple: el texto está en un solo run
    for run in parrafo.runs:
        if buscar in run.text:
            run.text = run.text.replace(buscar, reemplazar)
            return True

    # Caso complejo: el texto está partido entre varios runs
    # Reconstruimos el párrafo
    nuevo_texto = texto_completo.replace(buscar, reemplazar)
    if parrafo.runs:
        # Ponemos todo en el primer run y vaciamos los demás
        primer_run = parrafo.runs[0]
        primer_run.text = nuevo_texto
        for run in parrafo.runs[1:]:
            run.text = ""
        return True

    return False


def reemplazar_en_celda(celda, buscar: str, reemplazar: str) -> bool:
    """Reemplaza texto en todos los párrafos de una celda de tabla."""
    encontrado = False
    for para in celda.paragraphs:
        if reemplazar_en_parrafo(para, buscar, reemplazar):
            encontrado = True
    return encontrado


def reemplazar_en_documento(doc: Document, buscar: str, reemplazar: str) -> int:
    """
    Reemplaza todas las ocurrencias de 'buscar' por 'reemplazar' en todo el documento.
    Retorna el número de reemplazos realizados.
    """
    count = 0

    # En párrafos del cuerpo
    for para in doc.paragraphs:
        if reemplazar_en_parrafo(para, buscar, reemplazar):
            count += 1

    # En tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if reemplazar_en_celda(cell, buscar, reemplazar):
                    count += 1

    return count


def actualizar_celda_tabla(doc: Document, tabla_idx: int, fila_idx: int, col_idx: int, nuevo_texto: str):
    """Reemplaza el contenido completo de una celda específica."""
    tabla = doc.tables[tabla_idx]
    celda = tabla.rows[fila_idx].cells[col_idx]

    # Guardar formato del primer run del primer párrafo
    primer_para = celda.paragraphs[0] if celda.paragraphs else None

    # Limpiar la celda
    for para in celda.paragraphs:
        for run in para.runs:
            run.text = ""

    if primer_para and primer_para.runs:
        primer_para.runs[0].text = nuevo_texto
    elif primer_para:
        primer_para.clear()
        primer_para.add_run(nuevo_texto)


# ─────────────────────────────────────────────
# 3. CONSTRUCCIÓN DEL CONTRATO
# ─────────────────────────────────────────────

def aplicar_datos_al_contrato(doc: Document, datos: dict) -> Document:
    """
    Aplica los datos extraídos del Check_List al documento de contrato.
    """
    deudores = datos.get("deudores", [])
    acreedores = datos.get("acreedores", [])
    inmueble = datos.get("inmueble", {})
    prestamo = datos.get("prestamo", {})
    fecha = datos.get("fecha_firma", "")

    deudor = deudores[0] if deudores else {}
    acreedor1 = acreedores[0] if len(acreedores) > 0 else {}
    acreedor2 = acreedores[1] if len(acreedores) > 1 else {}

    # ── Datos del deudor ──
    reemplazos_deudor = [
        ("PATRICIA GONZÁLEZ VELÁSQUEZ", deudor.get("nombre_completo", "")),
        ("PATRICIA GONZ\u00c1LEZ VEL\u00c1SQUEZ", deudor.get("nombre_completo", "")),
        ("PATRICIA GONZALEZ VELASQUEZ", deudor.get("nombre_completo_mayuscula", "")),
        ("C.C. No. 52.202.940", f"C.C. No. {deudor.get('cc', '')}"),
        ("52.202.940 de Bogotá", f"{deudor.get('cc', '')} de {deudor.get('cc_expedicion', '')}"),
        ("52.202.940", deudor.get("cc", "")),
        ("Km 1.5 vía Cota Chia, casa 30, Cota", deudor.get("direccion", "")),
        ("Km 1,5 via Cota chia casa 30 Cota", deudor.get("direccion", "")),
        ("patricia.gonzalez.velasquez@hotmail.com", deudor.get("email", "")),
        ("+57 3005821543", deudor.get("telefono", "")),
        ("3005821543", deudor.get("telefono", "").replace("+57 ", "")),
        ("Casada con sociedad conyugal vigente", deudor.get("estado_civil", "")),
        ("mayor de edad y vecina del municipio de Cota", f"mayor de edad y vecina del municipio de {deudor.get('municipio', 'Cota')}"),
    ]

    # ── Datos Acreedor 1 ──
    reemplazos_acreedor1 = [
        ("HERNÁN JESÚS MONDOL CABARCAS", acreedor1.get("nombre_completo", "")),
        ("HERN\u00c1N JES\u00daS MONDOL CABARCAS", acreedor1.get("nombre_completo", "")),
        ("HERNAN JESUS MONDOL CABARCAS", acreedor1.get("nombre_completo_mayuscula", "")),
        ("C.C. No. 73.160.553", f"C.C. No. {acreedor1.get('cc', '')}"),
        ("73.160.553 de Cartagena", f"{acreedor1.get('cc', '')} de {acreedor1.get('cc_expedicion', '')}"),
        ("73.160.553", acreedor1.get("cc", "")),
        ("Cra 103b # 150C-30, casa 110, Agrupación Arboleda La Campiña, Bogotá D.C.", acreedor1.get("direccion", "")),
        ("hernanmondol73@gmail.com", acreedor1.get("email", "")),
        ("+ 57 3107668525", acreedor1.get("telefono", "")),
        ("+57 3107668525", acreedor1.get("telefono", "")),
        ("Soltero con unión marital de hecho", acreedor1.get("estado_civil", "")),
    ]

    # ── Datos Acreedor 2 ──
    reemplazos_acreedor2 = [
        ("HAWER ALBERTO HERRERA RODRÍGUEZ", acreedor2.get("nombre_completo", "")),
        ("HAWER ALBERTO HERRERA RODR\u00cdGUEZ", acreedor2.get("nombre_completo", "")),
        ("HAWER ALBERTO HERRERA RODRIGUEZ", acreedor2.get("nombre_completo_mayuscula", "")),
        ("C.C. No. 93.365.845", f"C.C. No. {acreedor2.get('cc', '')}"),
        ("93.365.845 de Ibagué", f"{acreedor2.get('cc', '')} de {acreedor2.get('cc_expedicion', '')}"),
        ("93.365.845", acreedor2.get("cc", "")),
        ("Carrera 6 # 53-29, Of. 604, Ibagué", acreedor2.get("direccion", "")),
        ("arqui.hawerh@gmail.com", acreedor2.get("email", "")),
        ("+57 3157909502", acreedor2.get("telefono", "")),
        ("Casado con sociedad conyugal vigente", acreedor2.get("estado_civil", "")),
    ]

    # ── Datos del inmueble ──
    reemplazos_inmueble = [
        ("50S-604333 de la Oficina de Registro de Instrumentos Públicos de Bogotá, Zona Sur.",
         f"{inmueble.get('matricula_inmobiliaria', '')} de la Oficina de Registro de Instrumentos Públicos de {inmueble.get('oficina_registro', '')}."),
        ("50S-604333", inmueble.get("matricula_inmobiliaria", "")),
        ("Calle 19 Sur No. 69A - 48, Bogotá D.C.", inmueble.get("direccion_corta", "")),
        ("AAA0041FAAF", inmueble.get("chip", "")),
    ]

    # ── Datos del préstamo ──
    monto_acr1_texto = acreedor1.get("participacion_texto", "")
    monto_acr1_formato = f"${acreedor1.get('participacion_monto', 0):,.0f}".replace(",", ".")
    cuota_ind_texto = acreedor1.get("cuota_mensual_texto", "")

    reemplazos_prestamo = [
        # Monto por acreedor (tabla datos básicos y tabla descripción)
        ("NOVENTA MILLONES DE PESOS MONEDA CORRIENTE (COP$90.000.000)", monto_acr1_texto),
        ("$90.000.000", monto_acr1_formato),
        ("90.000.000", f"{acreedor1.get('participacion_monto', 0):,.0f}".replace(",", ".")),
        # Tasa
        ("1.80% mensual anticipado", f"{prestamo.get('tasa_mensual', '')} mensual anticipado"),
        ("1.80%", prestamo.get("tasa_mensual", "")),
        # Plazo
        ("Sesenta (60) meses contados a partir de la Fecha de Desembolso del Monto Inicial",
         f"{prestamo.get('plazo_texto', '')} contados a partir de la Fecha de Desembolso del Monto Inicial"),
        # Cuota mensual individual
        ("Un millón seiscientos veinte mil pesos moneda corriente (COP$1.620.000)", cuota_ind_texto),
        ("$1.620.000", f"${acreedor1.get('cuota_mensual_individual', 0):,.0f}".replace(",", ".")),
        ("$1.710.000", f"${int(acreedor1.get('cuota_mensual_individual', 0) * 1.055555):,.0f}".replace(",", ".")),
        # Monto inicial (35%)
        ("TREINTA Y UN MILLONES QUINIENTOS MIL PESOS MONEDA CORRIENTE (COP$31.500.000)",
         prestamo.get("monto_inicial_texto", "")),
        # Monto restante (65%)
        ("CINCUENTA Y OCHO MILLONES QUINIENTOS MIL PESOS MONEDA CORRIENTE (COP$58.500.000)",
         prestamo.get("monto_restante_texto", "")),
        # Servicios Aluri
        ("$ 4.500.000", prestamo.get("servicios_aluri_texto", "")),
        # Fecha
        ("Ocho (8) de abril de dos mil veintiséis (2026)", fecha),
        ("ocho (8) de abril del año dos mil veintiséis (2026)", f"{fecha.lower()}"),
        ("ocho (8) del mes de abril del año dos mil veintiséis (2026)", f"{fecha.lower()}"),
        ("ocho (8) d\u00edas del mes de abril del a\u00f1o dos mil veintis\u00e9is (2026)", f"{fecha.lower()}"),
    ]

    # Aplicar todos los reemplazos
    todos_reemplazos = (
        reemplazos_deudor
        + reemplazos_acreedor1
        + reemplazos_acreedor2
        + reemplazos_inmueble
        + reemplazos_prestamo
    )

    total = 0
    for buscar, reemplazar in todos_reemplazos:
        if buscar and reemplazar and buscar != reemplazar:
            n = reemplazar_en_documento(doc, buscar, reemplazar)
            if n > 0:
                total += n
                print(f"  ✓ Reemplazado ({n}x): '{buscar[:60]}...' → '{reemplazar[:60]}...'")

    print(f"\n  Total de reemplazos realizados: {total}")
    return doc


# ─────────────────────────────────────────────
# 4. FLUJO PRINCIPAL DEL AGENTE
# ─────────────────────────────────────────────

def generar_contrato(
    ruta_checklist: Path,
    ruta_template: Path,
    ruta_output: Path
) -> Path:
    """
    Flujo principal:
    1. Leer Check_List
    2. Extraer datos con Claude
    3. Copiar template del contrato
    4. Aplicar datos al contrato
    5. Guardar contrato generado
    """

    print("\n" + "="*60)
    print("  AGENTE DE CONTRATOS - Contrato Mutuo Solo Interés")
    print("="*60)

    # Paso 1: Leer Check_List
    print(f"\n[1/4] Leyendo Check_List: {ruta_checklist.name}")
    texto = leer_texto_checklist(ruta_checklist)
    print(f"      → {len(texto.splitlines())} líneas extraídas")

    # Paso 2: Extraer datos con Claude
    print("\n[2/4] Extrayendo datos con Claude API...")
    datos = extraer_datos_con_claude(texto)

    deudor_nombre = datos.get("deudores", [{}])[0].get("nombre_completo_mayuscula", "DEUDOR")
    print(f"      → Deudor: {deudor_nombre}")
    print(f"      → Acreedores: {len(datos.get('acreedores', []))}")
    print(f"      → Monto: {datos.get('prestamo', {}).get('monto_total_texto', '')}")

    # Paso 3: Copiar template
    print(f"\n[3/4] Preparando documento base: {ruta_template.name}")
    ruta_output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(str(ruta_template), str(ruta_output))
    doc = Document(str(ruta_output))
    print(f"      → Copia creada en: {ruta_output}")

    # Paso 4: Aplicar datos
    print("\n[4/4] Aplicando datos al contrato...")
    doc = aplicar_datos_al_contrato(doc, datos)

    # Guardar
    doc.save(str(ruta_output))

    print(f"\n{'='*60}")
    print(f"  ✅ Contrato generado exitosamente:")
    print(f"     {ruta_output}")
    print(f"{'='*60}\n")

    # Guardar datos extraídos como JSON para referencia
    ruta_json = ruta_output.parent / (ruta_output.stem + "_datos.json")
    with open(str(ruta_json), "w", encoding="utf-8") as f:
        json.dump(datos, f, ensure_ascii=False, indent=2)
    print(f"  📋 Datos extraídos guardados en: {ruta_json.name}\n")

    return ruta_output


def main():
    parser = argparse.ArgumentParser(
        description="Genera Contrato Mutuo Solo Interés a partir de un Check_List"
    )
    parser.add_argument(
        "--checklist",
        type=Path,
        default=DEFAULT_CHECKLIST,
        help=f"Ruta al Check_List.docx (default: {DEFAULT_CHECKLIST})"
    )
    parser.add_argument(
        "--template",
        type=Path,
        default=DEFAULT_TEMPLATE,
        help=f"Ruta al template del contrato (default: {DEFAULT_TEMPLATE})"
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=None,
        help="Ruta del archivo de salida (default: Contratos_Generados/Contrato_[DEUDOR]_[fecha].docx)"
    )
    args = parser.parse_args()

    # Validar inputs
    if not args.checklist.exists():
        print(f"❌ Error: No se encontró el Check_List en: {args.checklist}")
        sys.exit(1)
    if not args.template.exists():
        print(f"❌ Error: No se encontró el template del contrato en: {args.template}")
        sys.exit(1)

    # Determinar ruta de salida
    if args.output:
        ruta_output = args.output
    else:
        hoy = date.today().strftime("%Y%m%d")
        ruta_output = DEFAULT_OUTPUT / f"Contrato_Mutuo_Solo_Interes_{hoy}.docx"

    try:
        generar_contrato(args.checklist, args.template, ruta_output)
    except json.JSONDecodeError as e:
        print(f"\n❌ Error al parsear respuesta de Claude: {e}")
        traceback.print_exc()
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ Error inesperado: {e}")
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
